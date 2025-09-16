# app.py
import streamlit as st
import requests
import subprocess
import sys
from datetime import datetime
from io import BytesIO
import urllib.parse

# -------------------------------
# Ensure python-docx is available (install only if missing)
# -------------------------------
try:
    from docx import Document
except Exception:
    # Try to install python-docx at runtime (falls back if environment allows)
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx==0.8.11"])
        from docx import Document
    except Exception:
        Document = None  # We'll still allow the app to run; docx export will show an error later.

# -------------------------------
# Config & Styling
# -------------------------------
st.set_page_config(page_title="InMind", layout="wide")

ACCENT = "#FDD2DC"

st.markdown(
    """
    <style>
    .stApp {{
        background-color: #000000;
        color: #ffffff;
        margin: 0;
        padding: 0;
    }}
    [data-testid="stAppViewContainer"] {{
        padding: 0;
        margin: 0;
    }}
    [data-testid="stHeader"] {{
        display: none;
    }}
    h1, h2, h3 {{
        color: #ffffff;
        text-align: center;
        font-family: 'Helvetica Neue', sans-serif;
    }}
    a, .accent {{ color: {ACCENT}; }}
    .stChatMessage {{ font-size: 1.05em; line-height: 1.6; }}
    .footer {{ font-size: 0.8em; text-align: center; color: #888888; margin-top: 3em; }}
    div[data-baseweb="base-input"] > div {{
        border-color: {ACCENT} !important;
    }}
    button[kind="primary"] {{
        background: {ACCENT} !important;
        color: #000 !important;
        border-radius: 8px !important;
    }}
    /* ensure chat area stretches */
    .css-1d391kg {{}}
    </style>
    """.format(ACCENT=ACCENT),
    unsafe_allow_html=True
)

# -------------------------------
# Branding (logo + subtitle)
# Keep the logo size and placement consistent with your request.
# -------------------------------
LOGO_FILENAME = "LOGO_PATH.png"  # keep this name or update to your file name in the repo
GITHUB_RAW = f"https://raw.githubusercontent.com/libbymatheus/InMind/main/{LOGO_FILENAME}"

st.markdown(
    f"""
    <div style="display:flex; justify-content:center; align-items:center; margin-bottom:12px;">
        <img src="{GITHUB_RAW}" alt="InMind Logo" width="200" style="display:block;"/>
    </div>
    """,
    unsafe_allow_html=True
)

st.markdown(
    """
    <div style="text-align: center; font-size:18px;">
        A free, conversational assistant.
    </div>
    """,
    unsafe_allow_html=True
)

# -------------------------------
# Free responder: health rules + Wikipedia fallback
# -------------------------------

def query_wikipedia(query, max_chars=900):
    try:
        params = {
            "action": "query",
            "list": "search",
            "srsearch": query,
            "srlimit": 1,
            "format": "json",
        }
        r = requests.get("https://en.wikipedia.org/w/api.php", params=params, timeout=8)
        if r.status_code != 200:
            return None
        data = r.json()
        if not data.get("query") or not data["query"]["search"]:
            return None
        page_title = data["query"]["search"][0]["title"]

        params2 = {
            "action": "query",
            "prop": "extracts",
            "exintro": True,
            "explaintext": True,
            "titles": page_title,
            "format": "json",
        }
        r2 = requests.get("https://en.wikipedia.org/w/api.php", params=params2, timeout=8)
        if r2.status_code != 200:
            return None
        pages = r2.json().get("query", {}).get("pages", {})
        if not pages:
            return None
        page = next(iter(pages.values()))
        extract = page.get("extract", "")
        if not extract:
            return None
        if len(extract) > max_chars:
            extract = extract[:max_chars].rsplit(".", 1)[0] + "..."
        link = f"https://en.wikipedia.org/wiki/{urllib.parse.quote(page_title)}"
        return f"**From Wikipedia: {page_title}**\n\n{extract}\n\n(See more: {link})"
    except Exception:
        return None

def health_response(text):
    t = text.lower()

    # Memory / Alzheimer's
    if any(k in t for k in ["memory", "forget", "dementia", "alzheimer", "confusion", "repeating"]):
        return (
            "**Possible area:** Possible Alzheimer’s-type cognitive decline (screening recommended).\n\n"
            "**Watch for:** progressive short-term memory loss, repeating questions, disorientation.\n\n"
            "**Next steps:** schedule a cognitive screen (MMSE/MoCA) with primary care or neurology; bring a family member to appointments; keep a symptom diary.\n\n"
            "**Meanwhile:** encourage regular exercise, sleep hygiene, cognitive activity, and social engagement.\n\n"
            "**This is not medical advice. Consult a clinician.**"
        )

    # Parkinson's / movement
    if any(k in t for k in ["tremor", "shake", "shaky", "stiff", "rigid", "parkinson", "bradykinesia", "slow movement"]):
        return (
            "**Possible area:** Movement disorder such as Parkinson’s disease.\n\n"
            "**Watch for:** tremor, slowed movements, small handwriting, balance problems.\n\n"
            "**Next steps:** request a neurological/movement disorder evaluation; record short videos of symptoms to show clinicians.\n\n"
            "**Meanwhile:** fall-safety planning, gentle exercise and stretching; do not stop medications without advice.\n\n"
            "**This is not medical advice. Consult a clinician.**"
        )

    # Stroke / FAST
    if any(k in t for k in ["stroke", "face", "arm", "speech", "face droop", "arm weakness", "slurred"]):
        return (
            "**Possible area:** Stroke (urgent if sudden).\n\n"
            "**FAST warning signs:** Face drooping, Arm weakness, Speech difficulty — if sudden call emergency services immediately.\n\n"
            "**Next steps:** emergency care for sudden symptoms; otherwise see neurology for evaluation.\n\n"
            "**This is not medical advice.**"
        )

    # Speech / language / FTD
    if any(k in t for k in ["speech", "words", "aphasia", "language", "word-finding"]):
        return (
            "**Possible area:** Language disorder (e.g., aphasia, primary progressive aphasia, or other neuro conditions).\n\n"
            "**Watch for:** progressive difficulty finding words, understanding, or forming sentences.\n\n"
            "**Next steps:** cognitive and speech-language evaluation; consider neuroimaging if advised.\n\n"
            "**This is not medical advice. Consult a clinician.**"
        )

    # Sleep disorders
    if any(k in t for k in ["sleep", "insomnia", "apnea", "daytime sleepiness", "restless legs", "narcole"]):
        return (
            "**Possible area:** Sleep disorder (insomnia, sleep apnea, restless legs, etc.).\n\n"
            "**Watch for:** loud snoring, gasping at night, excessive daytime sleepiness.\n\n"
            "**Next steps:** keep a sleep diary; consider referral to a sleep specialist or primary care evaluation.\n\n"
            "**Meanwhile:** maintain regular sleep schedule, reduce caffeine/alcohol before bed.\n\n"
            "**This is not medical advice.**"
        )

    # Mood: depression/anxiety
    if any(k in t for k in ["depress", "sad", "hopeless", "anxiety", "panic", "worried", "stress"]):
        return (
            "**Possible area:** Mood disorder (depression or anxiety).\n\n"
            "**Watch for:** persistent low mood, loss of interest, panic attacks, changes in sleep/appetite.\n\n"
            "**Next steps:** mental health screening with a clinician; consider therapy and support groups.\n\n"
            "**If suicidal thoughts occur, seek emergency help immediately.**"
        )

    # Seizure
    if any(k in t for k in ["seizure", "convulsion", "fit", "epilepsy", "blackout"]):
        return (
            "**Possible area:** Seizure disorder.\n\n"
            "**Watch for:** convulsions, loss of awareness, tongue biting, prolonged confusion after event.\n\n"
            "**Next steps:** urgent neurology evaluation and EEG if first-time event; seek emergency care for prolonged seizure.\n\n"
            "**This is not medical advice.**"
        )

    # Headache / migraine
    if any(k in t for k in ["headache", "migraine", "aura", "thunderclap", "sudden severe headache"]):
        return (
            "**Possible area:** Headaches or migraine.\n\n"
            "**Watch for:** sudden severe headache, neurological deficits — seek emergency care if present.\n\n"
            "**Next steps:** track triggers, consult primary care/neurology if frequent or severe.\n\n"
            "**This is not medical advice.**"
        )

    # Vision / eye
    if any(k in t for k in ["vision", "blur", "double vision", "blind", "sudden vision"]):
        return (
            "**Possible area:** Visual disturbance.\n\n"
            "**Watch for:** sudden vision loss or double vision — emergency if sudden.\n\n"
            "**Next steps:** ophthalmology or urgent care depending on onset.\n\n"
            "**This is not medical advice.**"
        )

    # Peripheral neuropathy / numbness
    if any(k in t for k in ["numb", "tingl", "pins and needles", "neuropathy", "burning in feet"]):
        return (
            "**Possible area:** Peripheral neuropathy.\n\n"
            "**Watch for:** progressive numbness, burning, weakness in limbs.\n\n"
            "**Next steps:** discuss with primary care; consider blood tests (glucose, B12) and neurology referral.\n\n"
            "**This is not medical advice.**"
        )

    return None

def generate_reply(user_text):
    # 1) health-first
    hr = health_response(user_text)
    if hr:
        return hr
    # 2) wikipedia fallback
    wiki = query_wikipedia(user_text)
    if wiki:
        return wiki
    # 3) generic fallback
    return (
        "I couldn't find a concise match. Try rephrasing with keywords (e.g., 'memory loss in older adult', "
        "'tremor and stiffness', or ask about a general topic)."
    )

# -------------------------------
# Conversation state (keeps history)
# -------------------------------
if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "assistant", "content": "Hi! I'm InMind. Ask me anything.", "time": datetime.now()}
    ]

# render chat history
for m in st.session_state.messages:
    with st.chat_message("assistant" if m["role"] == "assistant" else "user"):
        st.markdown(m["content"])

# -------------------------------
# Chat input
# -------------------------------
if prompt := st.chat_input("Type your question here..."):
    st.session_state.messages.append({"role": "user", "content": prompt, "time": datetime.now()})
    with st.chat_message("user"):
        st.markdown(prompt)

    # generate reply (free)
    reply = generate_reply(prompt)
    st.session_state.messages.append({"role": "assistant", "content": reply, "time": datetime.now()})
    with st.chat_message("assistant"):
        st.markdown(reply)

# -------------------------------
# Transcript export & clear
# -------------------------------
st.markdown("---")
st.markdown(
    "<div class='footer'>Disclaimer: InMind AI is for educational purposes only and does not provide medical diagnoses. Always consult a licensed healthcare professional for medical advice.</div>",
    unsafe_allow_html=True
)

if st.session_state.get("messages"):
    # TXT transcript
    transcript_txt = "InMind Chat Transcript\n"
    transcript_txt += f"Session started: {st.session_state.messages[0]['time'].strftime('%Y-%m-%d %H:%M:%S')}\n"
    transcript_txt += "-" * 40 + "\n\n"
    for m in st.session_state.messages:
        ts = m["time"].strftime("%H:%M:%S")
        speaker = "You" if m["role"] == "user" else "InMind"
        transcript_txt += f"[{ts}] {speaker}: {m['content']}\n\n"

    st.download_button(
        label="📥 Download Chat History (.txt)",
        data=transcript_txt,
        file_name="inmind_chat.txt",
        mime="text/plain"
    )

    # DOCX transcript (in-memory)
    if Document is not None:
        doc = Document()
        doc.add_heading("InMind Chat Transcript", level=1)
        doc.add_paragraph(f"Session started: {st.session_state.messages[0]['time'].strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("-" * 40)
        for m in st.session_state.messages:
            ts = m["time"].strftime("%H:%M:%S")
            speaker = "You" if m["role"] == "user" else "InMind"
            p = doc.add_paragraph()
            p.add_run(f"[{ts}] {speaker}: ").bold = True
            p.add_run(m["content"])
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button(
            label="📥 Download Chat History (.docx)",
            data=buf,
            file_name="inmind_chat.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.info("DOCX export is unavailable because python-docx could not be installed in this environment.")

# Clear chat button
if st.button("🗑️ Clear Chat"):
    st.session_state.messages = [
        {"role": "assistant", "content": "Hi! I'm InMind. Ask me anything.", "time": datetime.now()}
    ]
    st.experimental_rerun()
